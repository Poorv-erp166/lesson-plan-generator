import os
import pdfplumber
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxcompose.composer import Composer

genai.configure(api_key="YOUR_GEMINI_API_KEY")  # ← Replace this

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        print("❌ PDF Read Error:", e)
        return ""

def generate_lesson_plan(text, teacher_name, email, class_section):
    prompt = f"""
You are a trained school teacher. Based on the chapter content below, generate a detailed lesson plan in this format:

SANSKRUTI- AN ENGLISH MEDIUM SCHOOL  
LESSON PLAN (Chapter wise)  
SESSION-2025-2026

NAME OF THE TEACHER: [{teacher_name}]
EMAIL ID: [{email}]
PHONE NO.: [6264728390]
CLASS & SECTIONS: [{class_section}]
NAME OF THE CHAPTER/LESSON: 
OBJECTIVES:
LEARNING OUTCOMES:
DURATION:
PEDAGOGY TO BE USED:
RESOURCE REQUIRED:
ART INTEGRATION ACTIVITY:
SEA(S):
STUDENT ENGAGEMENT PLAN:
NO. OF HOME FUNS WITH DETAILS:
NO. OF WORKSHEET TO BE SHARED:
NO. OF PPTs TO BE SHARED: 0
COMPETENCY BASED QUESTIONS BASED ON BLOOM'S TAXONOMY:
NO. OF ASSESSMENTS:
PROJECTS IF ANY:
MULTIDISCIPLINARY ASPECTS:
SIGNATURE WITH DATE: {teacher_name} / 24/06/2025

CHAPTER TEXT:
{text[:4000]}
    """
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        print("🤖 Model to be used: gemini-1.5-flash")
        response = model.generate_content(prompt)
        return response.text if response.text else ""
    except Exception as e:
        print("❌ Gemini API Error:", e)
        return ""

def save_to_word(content, output_path="Sanskriti_Lesson_Plan.docx"):
    try:
        template_path = "templates/Sanskriti_Template.docx"
        template_doc = Document(template_path)
        composer = Composer(template_doc)

        new_doc = Document()
        new_doc.add_heading("SANSKRUTI-AN ENGLISH MEDIUM SCHOOL", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = new_doc.add_paragraph("SESSION: 2025–26")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        new_doc.add_paragraph("")

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                new_doc.add_paragraph("")
                continue
            if ":" in line:
                title, value = line.split(":", 1)
                p = new_doc.add_paragraph()
                run1 = p.add_run(f"{title.strip()}: ")
                run1.bold = True
                run1.font.size = Pt(11)
                run2 = p.add_run(value.strip())
                run2.font.size = Pt(11)
            else:
                new_doc.add_paragraph(line)

        temp_path = "temp_generated.docx"
        new_doc.save(temp_path)

        temp_doc = Document(temp_path)
        composer.append(temp_doc)
        composer.save(output_path)
        print(f"✅ Word file saved: {output_path}")
    except Exception as e:
        print("❌ Word Save Error:", e)
